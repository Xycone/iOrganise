from fastapi import HTTPException, UploadFile

import os
import gc

import torch
import tensorflow as tf
import tempfile
import aiofiles
import magic

import fitz
from typing import Optional
from docx import Document

from passlib.context import CryptContext
from datetime import datetime, timedelta
import jwt

import magic

def get_file_type(path):
    try:
        print(f"Checking file type for: {path}")  # Print the path of the file being checked
        
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(path)
        
        print(f"File type detected: {file_type}")  # Print the detected mime type
        return file_type
    except Exception as e:
        print(f"Error occurred while checking file type for {path}: {str(e)}")  # Print the error if one occurs
        return None

    
def is_video(path):
    file_type = get_file_type(path)
    return bool(file_type) and file_type.startswith("video/")

def is_audio(path):
    file_type = get_file_type(path)
    return bool(file_type) and file_type.startswith("audio/")

def is_image(path):
    file_type = get_file_type(path)
    return bool(file_type) and file_type.startswith("image/")

def is_text(path):
    return

def extract_text_from_pdf(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    text = []

    # Extract text from headers FIRST (they appear at the top)
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                text.append(para.text)

    # Extract main document text (paragraphs and tables in order)
    for para in doc.paragraphs:
        text.append(para.text)  # Extracting paragraph text

    # Extract table rows in order
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):  # Ensure there's text in the row
                text.append(" | ".join(row_text))

    # Extract footers LAST (they appear at the bottom)
    for section in doc.sections:
        if section.footer:
            for para in section.footer.paragraphs:
                text.append(para.text)

    return "\n".join(filter(None, text))  # Remove empty lines and join


def extract_text_from_txt(txt_path: str) -> str:
    with open(txt_path, 'r', encoding='utf-8') as file:
        return file.read()

# file uploading
async def save_uploaded_file(email: str, file: UploadFile):
    try:
        file_content = await file.read()

        with tempfile.NamedTemporaryFile(delete=True) as temp_file:
            temp_file.write(file_content)
            type = get_file_type(temp_file.name) if get_file_type(temp_file.name) else 'unknown'

        size = round(file.size / (1024 * 1024), 2)
        path = os.path.join("/app/file_storage", email, file.filename)
        
        async with aiofiles.open(path, 'wb') as output_file:
            await output_file.write(file_content)

        return type, size, path

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving file to storage: {str(e)}")

# authentication
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict, expires_delta: timedelta = None):
    secret_key = os.getenv('SECRET_KEY')
    if not secret_key:
        raise ValueError("SECRET_KEY environment variable not set.")
    
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=60))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, secret_key, algorithm="HS256")

def verify_jwt_token(token: str):
    try:
        payload = jwt.decode(token, os.getenv("SECRET_KEY"), algorithms="HS256")
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid JWT token")
        return user_id
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.JWTError:
        raise HTTPException(status_code=401, detail="Invalid JWT token")

# memory management
def free_memory():
    gc.collect()
    
    if torch.cuda.is_available():
        torch.cuda.empty_cache()